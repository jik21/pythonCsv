import os
import shutil
import csv
from docx import Document
from docx.shared import Inches


class DocRow:
    def __init__(self, csvRow):
        self.tableName = csvRow[0]
        self.colName = csvRow[1]
        self.dataType = csvRow[2]
        self.isMandate = not (csvRow[6] == 'Y')
        self.isPK = (csvRow[7] == 'Y')

def writeOutTable(doc, tableName, rows):
    if tableName == "" or not rows:
        return
    doc.add_heading('Table ' + tableName, level=2)
    doc.add_heading('Description of table ' + tableName, level=3)
    p = doc.add_paragraph(tableName + ' contains information of '+tableName+'.')
    doc.add_heading('Column list of the table ' + tableName, level=3)      
    table = initDocTable(doc,len(rows))
    for item in rows:
        row_cells = table.add_row().cells
        row_cells[0].text = item.colName
        row_cells[1].text = item.dataType
        row_cells[2].text = 'Y' if item.isMandate else 'N'
        row_cells[3].text = 'Y' if item.isPK else 'N'
    doc.add_page_break()    

def initDocTable(doc,rowNum):
    table = doc.add_table(rows=1, cols=4)
    headers = ['Data Item','Format','Mandatory','Primary Key']
    hdr_cells = table.rows[0].cells    
    for idx, name in enumerate(headers):
        paragraph = hdr_cells[idx].paragraphs[0]
        run = paragraph.add_run(name)
        run.bold = True
    #table.style = "Table Grid"
    return table
    
doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
with open('pcs_all_col_2.csv', newline='') as csvfile:                   
        spamreader = csv.reader(csvfile, delimiter=',', quotechar='"')        
        docRows = []
        processingTable = ""
        
        for row in spamreader:
            tableName = row[0]
            if tableName != processingTable:
                writeOutTable(doc,processingTable,docRows)
                processingTable = tableName
                docRows.clear()  
                
            docRows.append(DocRow(row))
                              
                #print (TableList)
        #off by 1    
        writeOutTable(doc,processingTable,docRows)        
        doc.save('pcs_all_col.docx')

        
