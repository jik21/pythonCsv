from docx import Document
from lxml import etree
from docxtpl import DocxTemplate, InlineImage
from docx.shared import Mm, Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import csv
import os

tpl=DocxTemplate('PSTemplate.docx')
imgDir = r'D:\tmp\rename'


def create_list(paragraph, list_type):
    p = paragraph._p #access to xml paragraph element
    pPr = p.get_or_add_pPr() #access paragraph properties
    numPr = OxmlElement('w:numPr') #create number properties element
    numId = OxmlElement('w:numId') #create numId element - sets bullet type
    numId.set(qn('w:val'), list_type) #set list type/indentation
    numPr.append(numId) #add bullet type to number properties list
    pPr.append(numPr) #add number properties to paragraph



def getCellText(table, row, cell):
    plist = table.rows[row].cells[cell].paragraphs
    return ' '.join([p.text for p in plist])

def getScreenFields(table, row, cell):
    fields = []
    for p in table.rows[row].cells[cell].paragraphs:
        if 'Action:' in p.text:
            break
        if isNumListItem(p):
            fields.append(p.text)
    return fields

def getActions(table, row, cell):
    actions = []
    reached = False
    for p in table.rows[row].cells[cell].paragraphs:
        if 'Action:' in p.text:
            reached = True            
        if 'Logic:' in p.text:
            break
        if isNumListItem(p) and reached:            
            actions.append(p.text)

    return actions

def getLogics(table, row, cell):
    logics = []
    reached = False
    for p in table.rows[row].cells[cell].paragraphs:
        if 'Logic:' in p.text:
            reached = True
        if isNumListItem(p) and reached:
            logics.append(p.text)
    return logics

def isNumListItem(p):
    return 'List' in p.style.name or r'<w:numPr>' in p._p.xml     

def getShapeInParagraph(plist):
    for p in plist:
        if 'w:drawing' in p._p.xml:
            x = etree.fromstring(p._p.xml)
            draw = x.findall('w:drawing',namespaces={
  'w': 'http://schemas.microsoft.com/office/word/2003/wordml',
  })
            s = InlineShape(draw)
            #draw = x.xpath('w:drawing',namespaces={
  #'w': 'http://schemas.microsoft.com/office/word/2003/wordml',
  #})
            print(etree.tostring(x, pretty_print=True))
        

def writeRowsToCsv(l):
    with open('eggs.csv', 'w', newline='', encoding='utf-8') as csvfile:
        spamwriter = csv.writer(csvfile, delimiter=',',
                                quotechar='|', quoting=csv.QUOTE_MINIMAL)
        for r in l:
            spamwriter.writerow(r)

def writeDocx(context, fileName):
    tpl.render(context)
    tpl.save(fileName)

def getActionDoc(doc, action):
    ordered = "5"
    unordered = "1"
    paragraph = doc.add_paragraph("On Page Load", "List Paragraph")
    create_list(paragraph, ordered)
    paragraph = doc.add_paragraph("Display", "List Paragraph")
    create_list(paragraph, unordered)
    for act in action:
        paragraph = doc.add_paragraph(act, "List Paragraph")
        create_list(paragraph, ordered)
        paragraph = doc.add_paragraph(" ", "List Paragraph")
        create_list(paragraph, unordered)
    
    return doc


def appendToContext(context, funcID, funcName, funcDesc, mode, screenField, action, logic, screens):
    pngFileName = funcID.replace('F-', '', 1) + '.png'
    function = {'ProgramID':funcID.replace('F-', 'P-', 1),
                'mode':mode,
                'funcName':funcName,
                'funcDesc':funcDesc,
                'ScreenID':funcID.replace('F-', 'S-', 1),
                'fields': [],
                'actionSubDoc': getActionDoc(tpl.new_subdoc(), action),
                'screenImage' : InlineImage(tpl,os.path.join(imgDir, pngFileName),width=Mm(150))    
        }
    for field in screenField:
        function['fields'].append({'name':field})
    context['functions'].append(function)

def parseFunc(document):    
    context = { 'functions':[]}
    for table in document.tables:
        funcID = getCellText(table, 1, 1)       
        funcName = getCellText(table,2,1)
        funcDesc = getCellText(table,5,1)
        mode = getCellText(table,6,1)
        logic = getLogics(table,11,1)
        action = getActions(table,11,1)
        screenField = getScreenFields(table,11,1)
        #screens = getShapeInParagraph(table.rows[13].cells[1].paragraphs)
        screens = table.rows[13].cells[1].paragraphs
        #l.append([funcID, funcName, funcDesc, mode, screenField, action, logic, screens])
        appendToContext(context, funcID, funcName, funcDesc, mode, screenField, action, logic, screens)
        
    return context

def main_func():
    document = Document('ss.docx')
    funcs = parseFunc(document)    
    writeDocx(funcs, "ps.docx")
    
main_func()
